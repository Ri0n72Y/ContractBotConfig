from __future__ import annotations

import contextlib
import re
import uuid
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..config.settings import GeneratorSettings, SUPPORTED_RENDER_PROFILES


class RenderError(RuntimeError):
    pass


class DocxRenderService:
    """Deterministic renderer for the supported contract Markdown subset."""

    def __init__(self, settings: GeneratorSettings) -> None:
        self.settings = settings

    @staticmethod
    def normalize_filename(value: str, fallback: str) -> str:
        name = Path(str(value or "").strip()).name
        if not name:
            name = fallback
        name = re.sub(r'[\\/:*?"<>|\x00-\x1f]+', "_", name).strip(" .")
        if not name:
            name = fallback
        if not name.lower().endswith(".docx"):
            name += ".docx"
        stem = Path(name).stem.strip() or "contract"
        encoded = stem.encode("utf-8")
        if len(encoded) > 180:
            raw = encoded[:180]
            while raw:
                try:
                    stem = raw.decode("utf-8")
                    break
                except UnicodeDecodeError:
                    raw = raw[:-1]
            stem = stem.rstrip(" ._") or "contract"
        return stem + ".docx"

    @staticmethod
    def validate_render_profile(value: str) -> str:
        profile = str(value or "standard_contract").strip() or "standard_contract"
        if profile not in SUPPORTED_RENDER_PROFILES:
            raise RenderError(
                "unsupported render_profile: "
                + profile
                + "; supported="
                + ",".join(SUPPORTED_RENDER_PROFILES)
            )
        return profile

    @staticmethod
    def _set_run_font(
        run: Any,
        font_name: str,
        font_size: float,
        bold: bool = False,
    ) -> None:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)

    def _style_standard_contract(self, document: Any) -> None:
        section = document.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(self.settings.margin_cm)
        section.bottom_margin = Cm(self.settings.margin_cm)
        section.left_margin = Cm(self.settings.margin_cm)
        section.right_margin = Cm(self.settings.margin_cm)

        normal = document.styles["Normal"]
        normal.font.name = self.settings.body_font
        normal.font.size = Pt(self.settings.body_font_size)
        normal._element.get_or_add_rPr().rFonts.set(
            qn("w:eastAsia"),
            self.settings.body_font,
        )
        normal.paragraph_format.line_spacing = self.settings.line_spacing
        normal.paragraph_format.space_after = Pt(0)

        for level in range(1, 4):
            style = document.styles[f"Heading {level}"]
            style.font.name = self.settings.heading_font
            style._element.get_or_add_rPr().rFonts.set(
                qn("w:eastAsia"),
                self.settings.heading_font,
            )
            style.font.size = Pt(
                max(
                    12.0,
                    self.settings.heading_font_size - (level - 1) * 2,
                )
            )
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)

        self._add_page_number_footer(section)

    def _apply_render_profile(self, document: Any, profile: str) -> None:
        if profile == "standard_contract":
            self._style_standard_contract(document)
            return
        raise RenderError("unsupported render_profile")

    def _add_page_number_footer(self, section: Any) -> None:
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = max(9.0, self.settings.body_font_size - 2)
        run = paragraph.add_run("第 ")
        self._set_run_font(run, self.settings.body_font, size)
        self._append_field(run, "PAGE")
        run = paragraph.add_run(" 页 共 ")
        self._set_run_font(run, self.settings.body_font, size)
        self._append_field(run, "NUMPAGES")
        run = paragraph.add_run(" 页")
        self._set_run_font(run, self.settings.body_font, size)

    @staticmethod
    def _append_field(run: Any, instruction: str) -> None:
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instr, separate, end])

    def _add_inline_text(self, paragraph: Any, text: str) -> None:
        parts = re.split(r"(\*\*.+?\*\*)", text)
        for part in parts:
            if not part:
                continue
            bold = part.startswith("**") and part.endswith("**") and len(part) >= 4
            value = part[2:-2] if bold else part
            run = paragraph.add_run(value)
            self._set_run_font(
                run,
                self.settings.body_font,
                self.settings.body_font_size,
                bold=bold,
            )

    @staticmethod
    def _table_separator(line: str) -> bool:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        return bool(cells) and all(
            re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells
        )

    @staticmethod
    def _table_cells(line: str) -> list[str]:
        return [cell.strip() for cell in line.strip().strip("|").split("|")]

    def _add_table(self, document: Any, rows: list[list[str]]) -> None:
        if not rows:
            return
        width = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=width)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_index, row in enumerate(rows):
            for column_index in range(width):
                cell = table.cell(row_index, column_index)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                value = row[column_index] if column_index < len(row) else ""
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_inline_text(paragraph, value)
                for run in paragraph.runs:
                    self._set_run_font(
                        run,
                        self.settings.body_font,
                        max(9.0, self.settings.body_font_size - 1),
                        bold=row_index == 0 or bool(run.bold),
                    )

    def _render_markdown(self, document: Any, markdown: str) -> None:
        lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        index = 0
        first_heading = True
        while index < len(lines):
            stripped = lines[index].strip()

            if stripped == "<!-- pagebreak -->":
                document.add_page_break()
                index += 1
                continue

            if stripped.startswith("|") and index + 1 < len(lines):
                next_line = lines[index + 1].strip()
                if next_line.startswith("|") and self._table_separator(next_line):
                    rows = [self._table_cells(stripped)]
                    index += 2
                    while index < len(lines) and lines[index].strip().startswith("|"):
                        rows.append(self._table_cells(lines[index].strip()))
                        index += 1
                    self._add_table(document, rows)
                    continue

            heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
            if heading:
                level = len(heading.group(1))
                text = heading.group(2).strip()
                paragraph = document.add_paragraph(style=f"Heading {level}")
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if level == 1 and first_heading
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                run = paragraph.add_run(text)
                self._set_run_font(
                    run,
                    self.settings.heading_font,
                    max(
                        12.0,
                        self.settings.heading_font_size - (level - 1) * 2,
                    ),
                    bold=True,
                )
                if level == 1:
                    first_heading = False
                index += 1
                continue

            bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
            if bullet:
                paragraph = document.add_paragraph(style="List Bullet")
                self._add_inline_text(paragraph, bullet.group(1))
                index += 1
                continue

            numbered = re.match(r"^(\d+[.)、])\s*(.+)$", stripped)
            if numbered:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.first_line_indent = Cm(0)
                self._add_inline_text(
                    paragraph,
                    f"{numbered.group(1)} {numbered.group(2)}",
                )
                index += 1
                continue

            if not stripped:
                index += 1
                continue

            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
            paragraph.paragraph_format.line_spacing = self.settings.line_spacing
            self._add_inline_text(paragraph, stripped)
            index += 1

    def render(
        self,
        *,
        document_title: str,
        document_markdown: str,
        output_filename: str,
        render_profile: str,
    ) -> dict[str, Any]:
        profile = self.validate_render_profile(render_profile)
        self.settings.output_dir.mkdir(parents=True, exist_ok=True)
        document = Document()
        self._apply_render_profile(document, profile)
        self._render_markdown(document, document_markdown)

        filename = self.normalize_filename(
            output_filename,
            fallback=self.normalize_filename(document_title, "contract.docx"),
        )
        unique_name = f"{uuid.uuid4().hex}_{filename}"
        output_path = (self.settings.output_dir / unique_name).resolve()
        if output_path.parent != self.settings.output_dir:
            raise RenderError("unsafe output path")
        try:
            document.save(output_path)
            size_bytes = output_path.stat().st_size
        except OSError as exc:
            with contextlib.suppress(OSError):
                output_path.unlink()
            raise RenderError("unable to save generated DOCX") from exc
        if size_bytes <= 0 or size_bytes > self.settings.max_file_bytes:
            with contextlib.suppress(OSError):
                output_path.unlink()
            raise RenderError("generated DOCX size is outside the allowed range")
        return {
            "success": True,
            "status": "ready",
            "output_path": str(output_path),
            "output_filename": filename,
            "size_bytes": size_bytes,
            "render_profile": profile,
        }
